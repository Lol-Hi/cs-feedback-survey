#Importing my other python files
from question_classes import NumericQn, CategoricalQn, FreeResponseQn, DemographicQn #python file containing classes for each question type, with functions to evaluate them

#Importing other libraries
import docx #allows me to write to a Microsoft Word Document

def console_output(question, exclude=[ ], summLen=5):
    """
    Prints the evaluation of the data to the console

    >>> console_output(DemographicQn(1, "Your Name"))
    
    ------------------
    
    Q1: Your Name – DemographicQn
    >>> console_output(NumericQn(2, "I would like to own a Microbit set for my own learning", [4, 4, 5, 1, 2, 2, 5, 5, 3]))
    
    ------------------
    
    Q2: I would like to own a Microbit set for my own learning – NumericQn
    Most popular choice: [5]
    
    Choices: 
    - 1 – 1 (11.1%)
    - 2 – 2 (22.2%)
    - 3 – 1 (11.1%)
    - 4 – 2 (22.2%)
    - 5 – 3 (33.3%)
    File name of pie chart: Q2_pie.png
    
    Average: 3.4444444444444446
    Median: 4
    >>> console_output(CategoricalQn(3, "I would consider using Microbit for my future school projects", ["Maybe", "Maybe", "Maybe", "Yes", "Yes", "Maybe", "Maybe", "Yes", "Maybe", "Maybe"]))
    
    ------------------
    
    Q3: I would consider using Microbit for my future school projects – CategoricalQn
    Most popular choice: ['Maybe']
    
    Choices: 
    - Maybe – 7 (70.0%)
    - Yes – 3 (30.0%)
    File name of pie chart: Q3_pie.png
    >>> console_output(FreeResponseQn(4, "What is your favourite part of the course?", ["shooting game", "Learning about the shooting game", "Everything", "Learning how to use the game block", "The games! :PPP", "Everything.", "Programming", "The coding", "The use of microbit to play the game I created"]))
    
    ------------------
    
    Q4: What is your favourite part of the course? – FreeResponseQn
    Summary of responses:
    - Everything.
    - Everything.
    - Learning about the shooting game.
    - Learning how to use the game block.
    - shooting game.
    """
    qType = question.__class__.__name__ #obtains the question type of the current question
    print("\n------------------\n") #prints a separator
    print("Q{}: {} – {}".format(question.qNumber, question.qStatement, qType)) #prints the question number, question statement and question type
    if qType != "DemographicQn": #if the question is a demographic question, there is no need to print anything else and the rest of the code below is skipped
        if qType == "FreeResponseQn": #if the question is a free response question, print the summary
            print("Summary of responses:")
            for sentence in question.summarize(sentenceNo=summLen, leaveOut=exclude): #iterates through each sentence in the summary, printing them out in point form
                print("- {}".format(sentence))
        else: #if the question is a numerical or categorical question, print the mode, list out the choices and create the pie chart
            print("Most popular choice: {}".format(question.mode())) #print the mode
            print("\nChoices: ")
            for choice, number in question.frequency().items(): #iterate through the choices and print out the number of responses per choice, as well as the percentage of responses per choice, in point form
                percent = question.freqPercent()[choice]
                print("- {} – {} ({:.1f}%)".format(choice, number, percent))
            print("File name of pie chart: {}".format(question.plot_pie())) #creates a pie chart and prints out the name of the file
            if qType == "NumericQn": #if the question is a numeric question, print the mean and the median
                print("\nAverage: {}".format(question.mean())) #print the mean
                print("Median: {}".format(question.median())) #print the median

def docx_output(question, doc, exclude=[ ], summLen=5):
    """
    Writes the evaluation of the data to a Microsoft Word document
    """
    qType = question.__class__.__name__ #obtains the question type of the current question
    doc.add_heading("Q{}: {} – {}".format(question.qNumber, question.qStatement, qType), 2) #writes the question number, question statement and question type as the heading of the analysis
    if qType != "DemographicQn": #if the question is a demographic question, nothing more has to be written and the rest of the code below is skipped
        if qType == "FreeResponseQn": #if the question is a free response question, output the summary
            summaryPara = doc.add_paragraph("Summary of responses: ")
            summaryPara.runs[0].bold = True #bold "Summary of responses: "
            for sentence in question.summarize(sentenceNo=summLen, leaveOut=exclude): #iterates throgh each sentence in the summary, writing them out in point form
                doc.add_paragraph(sentence, style="List Bullet")
        else: #if the question is a numerical or categorical question, output the mode, list out the choices and display the pie chart
            modePara = doc.add_paragraph("\nMost popular choice: ")
            modePara.add_run(str(question.mode())) #write the mode to the Word document
            modePara.runs[0].bold = True #bold "Most popular choice: "
            choicesPara = doc.add_paragraph("Choices: ")
            choicesPara.runs[0].bold = True #bold "Choices: "
            for choice, number in question.frequency().items(): #iterate through the choices and list out the number of responses per choice, as well as the percentage of responses per choice, in point form
                percent = question.freqPercent()[choice]
                doc.add_paragraph("{} – {} ({:.1f}%)".format(choice, number, percent), style="List Bullet")
            doc.add_picture(question.plot_pie(), height=docx.shared.Cm(8)) #add the pie chart to the Word document
            if qType == "NumericQn": #if the question is a numeric question. output the mean and median
                meanPara = doc.add_paragraph("Mean: ")
                meanPara.add_run(str(question.mean())) #write out the mean to the Word document
                meanPara.runs[0].bold = True #bold "Mean: "
                medianPara = doc.add_paragraph("Median: ")
                medianPara.add_run(str(question.median())) #write out the median to the Word document
                medianPara.runs[0].bold = True #bold "Median"
        doc.add_page_break() #for numerical, categorical and free response questions, add a page break as the analysis for them is too long to fit another question in the same page
            
